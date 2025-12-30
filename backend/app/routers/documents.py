from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import Response
from sqlalchemy.orm import Session
from pathlib import Path
from io import BytesIO
import tempfile
import os
from datetime import datetime
from typing import Optional
from app.database import get_db
from app import models
import json
import copy

router = APIRouter()

# Disport restrictions per destination - fixed text for each port
DISPORT_RESTRICTIONS = {
    "Shell Haven": """All vessels must be capable of connecting to two 16-inch Woodfield loading/unloading arms.
All vessels must be capable of discharging at a rate of 2500 Cubic meters per hour, or of maintaining a discharge pressure at the vessel's manifold of at least 100PSIG (7.5Bar).
It is Seller's responsibility to provide vessels which do not exceed the Maximum Limitations as follows: -
Maximum draft on arrival at S Jetty is 14.9 meters.
Max. LOA: 250 M
Max displacement of 135,000 MT
SDWT maximum 116,000 MT""",
    "Milford Haven": """All vessels must be capable of connecting to standard loading/unloading arms.
Maximum draft on arrival is 14.5 meters.
Max. LOA: 274 M
Max displacement of 150,000 MT
SDWT maximum 125,000 MT""",
    "Rotterdam": """All vessels must be capable of connecting to standard loading/unloading arms.
Maximum draft on arrival is 15.2 meters.
Max. LOA: 280 M
Max displacement of 160,000 MT
SDWT maximum 130,000 MT""",
    "Le Havre": """All vessels must be capable of connecting to standard loading/unloading arms.
Maximum draft on arrival is 14.0 meters.
Max. LOA: 260 M
Max displacement of 140,000 MT
SDWT maximum 115,000 MT""",
    "Naples": """All vessels must be capable of connecting to standard loading/unloading arms.
Maximum draft on arrival is 13.5 meters.
Max. LOA: 245 M
Max displacement of 130,000 MT
SDWT maximum 110,000 MT""",
}

def get_sheet_name_for_product(product_name: str, customer_name: str) -> str:
    """Determine which sheet to use based on product and customer"""
    product_lower = product_name.lower()
    customer_lower = customer_name.lower()
    
    # Check for KPCT customer
    if 'kpct' in customer_lower or 'kpc trading' in customer_lower:
        if 'jet' in product_lower or 'gas' in product_lower:
            return 'KPCT Gas-Jet'
        else:
            return 'KPCT Gas-Jet'  # Default for KPCT
    
    # Check product type
    if 'jet' in product_lower or 'gas' in product_lower:
        return 'Gas-Jet'
    elif 'fuel' in product_lower or 'hfo' in product_lower or 'heavy fuel' in product_lower:
        return 'Fuel'
    else:
        # Default to Gas-Jet
        return 'Gas-Jet'

def get_reference_data(db: Session, contract_number: str, customer_name: str):
    """Get reference data from the Reference sheet"""
    template_path = Path(__file__).parent.parent.parent / "templates" / "nomination_template.xlsx"
    from openpyxl import load_workbook
    
    wb = load_workbook(template_path, data_only=True)
    if 'Reference' not in wb.sheetnames:
        return None, None, None
    
    ws_ref = wb['Reference']
    
    # Search for matching contract/customer
    for row in range(2, ws_ref.max_row + 1):
        ref_contract = ws_ref.cell(row, 1).value or ''  # Column A: Contract type
        ref_name = ws_ref.cell(row, 2).value or ''  # Column B: Reference
        ref_full_name = ws_ref.cell(row, 3).value or ''  # Column C: Full name
        ref_md = ws_ref.cell(row, 4).value or ''  # Column D: MD reference
        ref_indicator = ws_ref.cell(row, 10).value or ''  # Column J: Indicator
        
        # Match by contract number or customer name
        if (contract_number and contract_number.lower() in str(ref_name).lower()) or \
           (customer_name and customer_name.lower() in str(ref_full_name).lower()):
            return ref_md, ref_indicator, ref_full_name
    
    return None, None, None

def get_inspector_info(db: Session, customer_name: str, inspector_name: str):
    """Get inspector info from LIST sheet"""
    template_path = Path(__file__).parent.parent.parent / "templates" / "nomination_template.xlsx"
    from openpyxl import load_workbook
    
    wb = load_workbook(template_path, data_only=True)
    if 'LIST ' not in wb.sheetnames:
        return None, None
    
    ws_list = wb['LIST ']
    
    # Search for customer in LIST sheet
    for row in range(2, ws_list.max_row + 1):
        list_customer = ws_list.cell(row, 3).value or ''  # Column C: Customer List
        list_inspector = ws_list.cell(row, 8).value or ''  # Column H: Inspector
        list_code = ws_list.cell(row, 9).value or ''  # Column I: Code
        list_spec = ws_list.cell(row, 10).value or ''  # Column J: Spec
        
        if customer_name and customer_name.lower() in str(list_customer).lower():
            # If inspector matches, return code and spec
            if inspector_name and inspector_name.lower() in str(list_inspector).lower():
                return list_code, list_spec
            # Otherwise return first match
            if not inspector_name:
                return list_code, list_spec
    
    return None, None

@router.get("/cargos/{cargo_id}/nomination")
def generate_nomination_excel(cargo_id: int, db: Session = Depends(get_db)):
    """Generate nomination Excel file for a specific cargo"""
    # Get cargo data
    cargo = db.query(models.Cargo).filter(models.Cargo.id == cargo_id).first()
    if not cargo:
        raise HTTPException(status_code=404, detail="Cargo not found")
    
    # Get contract
    contract = db.query(models.Contract).filter(models.Contract.id == cargo.contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")
    
    # Get customer
    customer = db.query(models.Customer).filter(models.Customer.id == cargo.customer_id).first()
    if not customer:
        raise HTTPException(status_code=404, detail="Customer not found")
    
    # Get monthly plan for laycan info
    monthly_plan = db.query(models.MonthlyPlan).filter(models.MonthlyPlan.id == cargo.monthly_plan_id).first()
    
    # Load Excel template
    template_path = Path(__file__).parent.parent.parent / "templates" / "nomination_template.xlsx"
    if not template_path.exists():
        raise HTTPException(status_code=500, detail=f"Nomination template not found at {template_path}")
    
    from openpyxl import load_workbook
    # Load template directly - keep original structure
    # Use keep_vba=False to avoid file corruption issues
    wb = load_workbook(template_path, read_only=False, keep_vba=False, data_only=False)
    
    # Remove external links to avoid security warnings
    # Clear external references if they exist
    if hasattr(wb, 'external_references'):
        wb.external_references = []
    if hasattr(wb, '_external_links'):
        wb._external_links = []
    
    # Determine which sheet to use
    sheet_name = get_sheet_name_for_product(cargo.product_name, customer.name)
    if sheet_name not in wb.sheetnames:
        # Fallback to first available template sheet
        available_sheets = [s for s in wb.sheetnames if s not in ['Reference', 'LIST ']]
        if not available_sheets:
            raise HTTPException(status_code=500, detail="No template sheet found")
        sheet_name = available_sheets[0]
    
    ws = wb[sheet_name]
    
    # Get reference data
    md_reference, indicator, ref_full_name = get_reference_data(db, contract.contract_number, customer.name)
    
    # Get inspector info
    inspector_code, spec_code = get_inspector_info(db, customer.name, cargo.inspector_name or '')
    
    # Fill in the template
    # Note: A4 (Date/Header) is kept as is in template - don't change it
    
    # ITEM (Row 22, Column C) - put red text "Enter Item Number"
    from openpyxl.styles import Font
    cell_c22 = ws['C22']
    cell_c22.value = 'Enter Item Number'
    if cell_c22.font is None:
        cell_c22.font = Font(color='FF0000')
    else:
        cell_c22.font = Font(color='FF0000', size=cell_c22.font.size, name=cell_c22.font.name)
    
    # LAYDAYS (Row 22, Columns G-H)
    laycan = cargo.laycan_window or ''
    if monthly_plan:
        if contract.contract_type.value == 'FOB':
            laycan = monthly_plan.laycan_2_days or monthly_plan.laycan_5_days or laycan
        else:
            laycan = cargo.laycan_window or laycan
    
    if laycan and laycan != 'TBA' and laycan != '-':
        # Try to parse laycan (e.g., "21-22/11" or "20-21/11/2025")
        import re
        match = re.search(r'(\d{1,2})-(\d{1,2})/(\d{1,2})', laycan)
        if match:
            day1, day2, month = match.groups()
            ws['G22'] = f'({day1}-{day2}/{month})'
        else:
            ws['G22'] = laycan
    else:
        ws['G22'] = 'TBA'
    
    # H22 - keep empty (no text)
    ws['H22'] = ''
    
    # VESSEL (Row 23, Column C)
    ws['C23'] = cargo.vessel_name
    
    # ETA KUWAIT (Row 23, Column G)
    eta = cargo.eta or 'TBA'
    ws['G23'] = eta
    
    # PRODUCT (Row 24, Column C)
    ws['C24'] = cargo.product_name
    
    # Delivery (Row 24, Column G) - FOB or CIF
    ws['G24'] = contract.contract_type.value
    
    # QUANTITY (Row 25, Column C) - remove .0 if whole number
    quantity = cargo.cargo_quantity
    if quantity == int(quantity):
        quantity_str = f"{int(quantity)} KT +/- 10%"
    else:
        quantity_str = f"{quantity} KT +/- 10%"
    ws['C25'] = quantity_str
    
    # Parse load ports once for reuse
    load_ports = cargo.load_ports or ''
    ports_list = []
    if load_ports:
        # Handle both comma and slash separators
        if ',' in load_ports:
            ports_list = [p.strip() for p in load_ports.split(',')]
        elif '/' in load_ports:
            ports_list = [p.strip() for p in load_ports.split('/')]
        else:
            ports_list = [load_ports.strip()]
    
    # Load ports in G7 to G14 (header section) - fill these with load ports
    # Note: G8 is kept as "KWT" - skip it
    # G9, G11, G13 should be the same as G7 (first port)
    if ports_list:
        first_port = ports_list[0] if ports_list else ''
        
        # G7 gets first port
        ws['G7'] = first_port
        
        # G9, G11, G13 get the same as G7 (first port)
        ws['G9'] = first_port
        ws['G11'] = first_port
        ws['G13'] = first_port
        
        # G10, G12, G14 get subsequent ports if available
        if len(ports_list) > 1:
            ws['G10'] = ports_list[1] if len(ports_list) > 1 else ''
        if len(ports_list) > 2:
            ws['G12'] = ports_list[2] if len(ports_list) > 2 else ''
        if len(ports_list) > 3:
            ws['G14'] = ports_list[3] if len(ports_list) > 3 else ''
    
    # Load ports info (Row 25, Column D and Row 26, Column C)
    if ports_list:
        if len(ports_list) == 1:
            ws['D25'] = f'(AS PER MASTER REQUEST) to be fully loaded Ex. {ports_list[0]}'
        elif len(ports_list) == 2:
            ws['D25'] = f'(AS PER MASTER REQUEST) to be Loaded as follows: '
            ws['C26'] = f' 1ST PORT: {ports_list[0]}\n 2ND PORT: Balance Quantity Ex. {ports_list[1]}'
        else:
            ws['D25'] = f'(AS PER MASTER REQUEST) to be loaded from: {load_ports}'
    
    # INSPECTOR (Row 28, Column C)
    inspector = cargo.inspector_name or 'TBA'
    ws['C28'] = inspector
    
    # CHARGES (Row 28, Column G-H) - Split customer name
    customer_name_parts = customer.name.split()
    if len(customer_name_parts) > 0:
        ws['G28'] = '50/50 KPC/ '
        ws['H28'] = customer_name_parts[0] if len(customer_name_parts) > 0 else customer.name
    else:
        ws['G28'] = '50/50 KPC/ '
        ws['H28'] = customer.name
    
    # SAMPLING KPC PROCEDURE (Row 29, Column C) - based on product type
    product_lower = cargo.product_name.lower().strip()
    # Check for gasoil (including "gasoil 10ppm" or "gasoil 10 ppm" variations)
    if 'gasoil' in product_lower:
        ws['C29'] = 'G-001'
    # Check for JET-A1
    elif 'jet' in product_lower and ('a1' in product_lower or 'a-1' in product_lower):
        ws['C29'] = 'K-001'
    else:
        # Default fallback to K-001
        ws['C29'] = 'K-001'
    
    # SPEC (Row 29, Column G) - from reference or inspector info
    if spec_code:
        ws['G29'] = spec_code
    else:
        # Default based on product
        product_lower = cargo.product_name.lower()
        if 'jet' in product_lower or 'gas' in product_lower:
            ws['G29'] = '3001-A'
        elif 'fuel' in product_lower or 'hfo' in product_lower:
            ws['G29'] = '5325 M'
        else:
            ws['G29'] = '3001-A'
    
    # CUSTOMER (Row 30, Column C)
    ws['C30'] = customer.name
    
    # CONTRACT NO. (Row 30, Column G)
    ws['G30'] = f'({contract.contract_number})'
    
    # Save workbook - use temporary file first to ensure proper Excel format
    temp_file_path = None
    try:
        # Ensure external links are removed before saving
        if hasattr(wb, 'external_references'):
            wb.external_references = []
        if hasattr(wb, '_external_links'):
            wb._external_links = []
        
        # Create temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx')
        temp_file_path = temp_file.name
        temp_file.close()
        
        # Save workbook to temporary file with write_only=False
        try:
            wb.save(temp_file_path)
        except Exception as save_error:
            wb.close()
            raise HTTPException(status_code=500, detail=f"Error saving Excel file: {str(save_error)}")
        
        wb.close()  # Close workbook immediately after saving
        
        # Verify file was created and has content
        if not os.path.exists(temp_file_path):
            raise HTTPException(status_code=500, detail="Failed to generate Excel file - file not created")
        
        file_size = os.path.getsize(temp_file_path)
        if file_size == 0:
            raise HTTPException(status_code=500, detail="Failed to generate Excel file - file is empty")
        
        # Validate it's a valid Excel file (check for ZIP signature - Excel files are ZIP archives)
        with open(temp_file_path, 'rb') as f:
            first_bytes = f.read(4)
            f.seek(0)  # Reset for later reading
            if first_bytes != b'PK\x03\x04':  # ZIP file signature
                raise HTTPException(status_code=500, detail="Generated file is not a valid Excel file")
        
        # Read the file content (file handle already open from validation, but we closed it)
        with open(temp_file_path, 'rb') as f:
            file_content = f.read()
        
        # Generate filename: vessel_name_customer_name_contract_number_load_port
        # Clean up names for filename (remove special characters)
        vessel_name_clean = (cargo.vessel_name or 'TBA').replace(' ', '_').replace('/', '_').replace('\\', '_').replace(':', '_')
        customer_name_clean = (customer.name or 'TBA').replace(' ', '_').replace('/', '_').replace('\\', '_').replace(':', '_')
        contract_number_clean = (contract.contract_number or 'TBA').replace(' ', '_').replace('/', '_').replace('\\', '_').replace(':', '_')
        # Get first load port or use TBA
        load_ports_str = cargo.load_ports or 'TBA'
        first_port = load_ports_str.split(',')[0].split('/')[0].strip() if load_ports_str != 'TBA' else 'TBA'
        load_port_clean = first_port.replace(' ', '_').replace('/', '_').replace('\\', '_').replace(':', '_')
        
        filename = f"{vessel_name_clean}_{customer_name_clean}_{contract_number_clean}_{load_port_clean}.xlsx"
        
        # Clean up temporary file
        try:
            if temp_file_path and os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
        except:
            pass
        
        # Return file as response
        # Use filename* for proper UTF-8 encoding support
        return Response(
            content=file_content,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
                "Content-Length": str(len(file_content))
            }
        )
    except Exception as e:
        # Ensure workbook is closed even on error
        try:
            wb.close()
        except:
            pass
        # Clean up temporary file on error
        try:
            if temp_file_path and os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
        except:
            pass
        raise HTTPException(status_code=500, detail=f"Error generating Excel file: {str(e)}")


@router.get("/tng/{monthly_plan_id}")
def generate_tng_document(
    monthly_plan_id: int, 
    format: str = Query("docx", description="Output format: docx or pdf"),
    db: Session = Depends(get_db)
):
    """Generate Tonnage Memo (TNG) document for a monthly plan (or combi group)"""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Get the monthly plan
    monthly_plan = db.query(models.MonthlyPlan).filter(models.MonthlyPlan.id == monthly_plan_id).first()
    if not monthly_plan:
        raise HTTPException(status_code=404, detail="Monthly plan not found")
    
    # Get all plans in combi group (if applicable)
    combi_plans = []
    if monthly_plan.combi_group_id:
        combi_plans = db.query(models.MonthlyPlan).filter(
            models.MonthlyPlan.combi_group_id == monthly_plan.combi_group_id
        ).all()
    else:
        combi_plans = [monthly_plan]
    
    # Get contract - try quarterly plan first, then direct contract_id
    contract = None
    if monthly_plan.quarterly_plan_id:
        quarterly_plan = db.query(models.QuarterlyPlan).filter(
            models.QuarterlyPlan.id == monthly_plan.quarterly_plan_id
        ).first()
        if quarterly_plan:
            contract = db.query(models.Contract).filter(
                models.Contract.id == quarterly_plan.contract_id
            ).first()
    
    if not contract and monthly_plan.contract_id:
        contract = db.query(models.Contract).filter(
            models.Contract.id == monthly_plan.contract_id
        ).first()
    
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found for this monthly plan")
    
    # Get customer
    customer = db.query(models.Customer).filter(models.Customer.id == contract.customer_id).first()
    if not customer:
        raise HTTPException(status_code=404, detail="Customer not found")
    
    # Get product names and quantities from combi plans
    products_data = []
    total_quantity = 0
    for plan in combi_plans:
        product_name = plan.product_name
        if not product_name and plan.quarterly_plan_id:
            qp = db.query(models.QuarterlyPlan).filter(models.QuarterlyPlan.id == plan.quarterly_plan_id).first()
            if qp:
                product_name = qp.product_name
        
        quantity = plan.month_quantity or 0
        total_quantity += quantity
        products_data.append({
            "name": product_name or "Unknown",
            "quantity": quantity
        })
    
    # Get discharge port and restrictions
    discharge_port = contract.cif_destination or "TBA"
    disport_restrictions = DISPORT_RESTRICTIONS.get(discharge_port, "Contact operations for vessel requirements.")
    
    # Load template
    template_path = Path(__file__).parent.parent.parent / "templates" / "tng_template.docx"
    if not template_path.exists():
        raise HTTPException(status_code=500, detail="TNG template not found")
    
    doc = Document(template_path)
    
    # Helper function to replace text in paragraphs
    def replace_in_paragraph(paragraph, old_text, new_text):
        if old_text in paragraph.text:
            # Preserve formatting by replacing in runs
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
            # If not found in runs, replace in full text (loses some formatting)
            if old_text in paragraph.text:
                paragraph.text = paragraph.text.replace(old_text, new_text)
    
    # Prepare replacement data
    today_date = datetime.now().strftime("%d %B %Y")
    loading_window = monthly_plan.loading_window or "TBA"
    delivery_window = monthly_plan.delivery_window or "TBA"
    
    # Build products table text
    # Format: DISCHARGE_PORT    Product1    50KT ± 10%    DELIVERY_WINDOW
    #                          Product2    50KT ± 10%
    products_text_lines = []
    for i, prod in enumerate(products_data):
        qty_str = f"{int(prod['quantity']) if prod['quantity'] == int(prod['quantity']) else prod['quantity']}KT ± 10%"
        if i == 0:
            # First product includes discharge port and delivery window
            products_text_lines.append(f"                {discharge_port.upper()}                   {prod['name']}                {qty_str}                  {delivery_window}")
        else:
            # Subsequent products only show product and quantity
            products_text_lines.append(f"                                                      {prod['name']}                {qty_str}")
    
    products_table_text = "\n".join(products_text_lines)
    
    # TNG notes from contract
    tng_notes = contract.tng_notes or """Cargo to be commingled.
Vessel to adopt early departure procedure (EDP). 
Master to send his daily ETA to the following email:  GXSTRMDCARGOJET@shell.com
BL WILL NOT BE AVAILABLE AT DISPORT."""
    
    # Get discharge ranges from contract
    discharge_ranges = contract.discharge_ranges or ""
    
    # Replace placeholders in document
    replacements = {
        "{{DATE}}": today_date,
        "{{CUSTOMER_NAME}}": customer.name,
        "{{CONTRACT_NUMBER}}": contract.contract_number,
        "{{LOADING_WINDOW}}": loading_window,
        "{{DELIVERY_WINDOW}}": delivery_window,
        "{{DISCHARGE_PORT}}": discharge_port.upper(),
        "{{PRODUCT_NAME}}": ", ".join([p["name"] for p in products_data]),
        "{{CARGO_QUANTITY}}": f"{int(total_quantity) if total_quantity == int(total_quantity) else total_quantity}KT ± 10%",
        "{{DISPORT_RESTRICTIONS}}": disport_restrictions,
        "{{DISCHARGE_RANGES}}": discharge_ranges,
        "{{TNG_NOTES}}": tng_notes,
    }
    
    # Process paragraphs
    for para in doc.paragraphs:
        for old_text, new_text in replacements.items():
            replace_in_paragraph(para, old_text, new_text)
        
        # Handle the products table line specially for combi cargos
        if "{{PRODUCT_NAME}}" in para.text or any(p["name"] in para.text for p in products_data):
            # If combi cargo with multiple products, format as multi-line
            if len(products_data) > 1:
                para.text = products_table_text
    
    # Process tables (for the SUB line with customer name and contract number)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        replace_in_paragraph(para, old_text, new_text)
    
    # Save to temporary file
    temp_file_path = None
    try:
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_file_path = temp_file.name
        temp_file.close()
        
        doc.save(temp_file_path)
        
        # Read file content
        with open(temp_file_path, 'rb') as f:
            file_content = f.read()
        
        # Generate filename
        customer_clean = customer.name.replace(' ', '_').replace('/', '_')[:20]
        contract_clean = contract.contract_number.replace(' ', '_').replace('/', '_')
        month_year = f"{monthly_plan.month}_{monthly_plan.year}"
        filename = f"TNG_{customer_clean}_{contract_clean}_{month_year}.docx"
        
        # Clean up temp file
        try:
            os.unlink(temp_file_path)
        except:
            pass
        
        return Response(
            content=file_content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
                "Content-Length": str(len(file_content))
            }
        )
        
    except Exception as e:
        # Clean up on error
        try:
            if temp_file_path and os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
        except:
            pass
        raise HTTPException(status_code=500, detail=f"Error generating TNG document: {str(e)}")

